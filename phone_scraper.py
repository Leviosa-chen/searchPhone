#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
网站手机号码和联系人爬虫
爬取指定网站的所有页面，提取手机号码和联系人信息
"""

import requests
from bs4 import BeautifulSoup
import re
import csv
import time
import urllib.parse
from urllib.parse import urljoin, urlparse
import logging
from typing import List, Dict, Set, Tuple
import json
from docx import Document
from docx.shared import Inches
import os

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('scraper.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class PhoneScraper:
    def __init__(self, base_url: str):
        self.base_url = base_url
        self.domain = urlparse(base_url).netloc
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        self.visited_urls: Set[str] = set()
        self.phone_contacts: List[Dict[str, str]] = []
        # 用于跟踪已出现的手机号和联系人，确保不重复
        self.seen_phones: Set[str] = set()
        self.seen_contacts: Set[str] = set()
        self.site_title = "未知网站"  # 网站标题
        # 可选的进度回调：接受 dict 参数
        self.progress_callback = None

    def _report(self, event_type: str, data: Dict = None) -> None:
        """向外部报告进度（如果已设置回调）。"""
        if not hasattr(self, 'progress_callback') or not callable(self.progress_callback):
            return
        payload = {'type': event_type}
        if data:
            payload.update(data)
        try:
            self.progress_callback(payload)
        except Exception:
            # 回调失败不影响主流程
            pass
        
    def clean_text(self, text: str) -> str:
        """清理文本，去掉特殊字符"""
        if not text:
            return ""
        
        # 去掉特殊字符，保留中文、英文、数字、空格、冒号、括号
        cleaned = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s：:()（）]', '', text)
        # 去掉多余空格
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned
    
    def is_valid_url(self, url: str) -> bool:
        """检查URL是否有效且属于同一域名"""
        try:
            parsed = urlparse(url)
            return parsed.netloc == self.domain and parsed.scheme in ['http', 'https']
        except:
            return False
    
    def extract_phone_numbers(self, text: str) -> List[str]:
        """提取文本中的手机号码 - 优化版，确保前后不是数字"""
        # 中国手机号码正则表达式 - 使用负向前瞻和负向后瞻确保前后不是数字
        phone_patterns = [
            r'(?<!\d)1[3-9]\d{9}(?!\d)',  # 标准11位手机号，前后不能是数字
            r'(?<!\d)\+86\s*1[3-9]\d{9}(?!\d)',  # 带+86前缀
            r'(?<!\d)86\s*1[3-9]\d{9}(?!\d)',  # 带86前缀
            r'(?<!\d)1[3-9]\d{2}\s*\d{4}\s*\d{4}(?!\d)',  # 带空格的手机号
            r'(?<!\d)1[3-9]\d{2}-\d{4}-\d{4}(?!\d)',  # 带连字符的手机号
        ]
        
        phones = []
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                # 清理号码格式
                clean_phone = re.sub(r'[\s\+\-]', '', match)
                if len(clean_phone) == 11 and clean_phone.startswith('1'):
                    # 额外验证：确保不是文件名的一部分
                    if not self._is_filename_part(clean_phone, text):
                        phones.append(clean_phone)
        
        return list(set(phones))  # 去重
    
    def _is_filename_part(self, phone: str, text: str) -> bool:
        """检查手机号是否是文件名的一部分"""
        # 查找手机号在文本中的位置
        phone_pos = text.find(phone)
        if phone_pos == -1:
            return False
        
        # 检查前后字符，判断是否是文件名
        start_pos = max(0, phone_pos - 10)
        end_pos = min(len(text), phone_pos + 21)
        context = text[start_pos:end_pos]
        
        # 如果包含常见文件扩展名，可能是文件名
        file_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.pdf', '.doc', '.docx', '.xls', '.xlsx']
        for ext in file_extensions:
            if ext in context.lower():
                return True
        
        # 如果前后都是数字，可能是长数字的一部分
        if phone_pos > 0 and phone_pos + 11 < len(text):
            prev_char = text[phone_pos - 1]
            next_char = text[phone_pos + 11]
            if prev_char.isdigit() and next_char.isdigit():
                return True
        
        return False
    
    def extract_contacts(self, text: str) -> List[str]:
        """提取联系人信息 - 优化版，注意前后语义"""
        # 改进的联系人正则表达式，更注重语义
        contact_patterns = [
            # 标准格式：关键词: 联系人
            r'(?:联系人|负责人|经理|主管|主任|总监|姓名|名字)[：:]\s*([^\n\r]{1,15})',
            
            # 带职务的格式：联系人：张三 经理
            r'(?:联系人|负责人|经理|主管|主任|总监)[：:]\s*([^\n\r]{1,20})',
            
            # 表格格式：联系人 张三
            r'(?:联系人|负责人|经理|主管|主任|总监)\s+([^\n\r]{1,15})',
            
            # 姓名格式：姓名：张三
            r'姓名[：:]\s*([^\n\r]{1,15})',
        ]
        
        contacts = []
        for pattern in contact_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                contact = match.strip()
                if self._is_valid_contact(contact):
                    contacts.append(contact)
        
        return list(set(contacts))  # 去重
    
    def _is_valid_contact(self, contact: str) -> bool:
        """验证联系人信息是否有效"""
        if not contact:
            return False
        
        # 长度检查
        if len(contact) < 2 or len(contact) > 20:
            return False
        
        # 排除明显无效的内容
        invalid_patterns = [
            r'^\d+$',  # 纯数字
            r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$',  # 邮箱
            r'^1[3-9]\d{9}$',  # 手机号
            r'^\d{3,4}-\d{7,8}$',  # 座机号
            r'^[^\u4e00-\u9fa5a-zA-Z]*$',  # 不包含中文或英文
        ]
        
        for pattern in invalid_patterns:
            if re.match(pattern, contact):
                return False
        
        # 必须包含中文或英文
        if not re.search(r'[\u4e00-\u9fa5a-zA-Z]', contact):
            return False
        
        return True
    
    def get_page_content(self, url: str) -> Tuple[str, BeautifulSoup]:
        """获取页面内容"""
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            response.encoding = response.apparent_encoding or 'utf-8'
            
            soup = BeautifulSoup(response.text, 'html.parser')
            return response.text, soup
        except Exception as e:
            logger.error(f"获取页面失败 {url}: {e}")
            return "", None
    
    def find_all_links(self, soup: BeautifulSoup, current_url: str) -> List[str]:
        """查找页面中的所有链接"""
        links = []
        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href']
            absolute_url = urljoin(current_url, href)
            
            if self.is_valid_url(absolute_url):
                links.append(absolute_url)
        
        return links
    
    def extract_page_info(self, url: str, soup: BeautifulSoup) -> None:
        """提取页面中的手机号码和联系人信息"""
        if not soup:
            return
            
        # 获取页面文本内容
        text_content = soup.get_text()
        
        # 提取手机号码
        phones = self.extract_phone_numbers(text_content)
        
        # 提取联系人
        contacts = self.extract_contacts(text_content)
        
        # 去重处理：只保留首次出现的手机号和联系人
        unique_phones = []
        for phone in phones:
            if phone not in self.seen_phones:
                unique_phones.append(phone)
                self.seen_phones.add(phone)
        
        unique_contacts = []
        for contact in contacts:
            if contact not in self.seen_contacts:
                unique_contacts.append(contact)
                self.seen_contacts.add(contact)
        
        # 获取页面标题
        title = soup.find('title')
        page_title = self.clean_text(title.get_text().strip()) if title else "无标题"
        
        # 记录找到的信息
        if unique_phones or unique_contacts:
            page_info = {
                'url': url,
                'title': page_title,
                'phone_numbers': '; '.join(unique_phones),
                'contacts': '; '.join(unique_contacts),
                'phone_count': len(unique_phones),
                'contact_count': len(unique_contacts),
                'original_phones': '; '.join(phones),  # 保留原始数据用于对比
                'original_contacts': '; '.join(contacts)
            }
            self.phone_contacts.append(page_info)
            logger.info(f"页面 {url} 找到 {len(unique_phones)} 个新手机号, {len(unique_contacts)} 个新联系人")
            
            # 如果有重复数据，记录去重信息
            duplicate_phones = len(phones) - len(unique_phones)
            duplicate_contacts = len(contacts) - len(unique_contacts)
            if duplicate_phones > 0 or duplicate_contacts > 0:
                logger.info(f"去重: {duplicate_phones} 个重复手机号, {duplicate_contacts} 个重复联系人")
    
    def crawl_website(self, max_pages: int = None) -> None:
        """爬取网站。
        当 max_pages 为 None 时按安全上限爬取；否则最多爬取 max_pages 页。
        """
        urls_to_visit = [self.base_url]
        page_count = 0
        
        logger.info(f"开始爬取网站: {self.base_url}")
        if max_pages is not None:
            try:
                max_pages_int = int(max_pages)
                if max_pages_int > 0:
                    logger.info(f"已启用页数限制: 最多爬取 {max_pages_int} 页")
                else:
                    max_pages_int = None
            except Exception:
                max_pages_int = None
        else:
            max_pages_int = None
        if max_pages_int is None:
            logger.info("注意: 未设置页数限制，将按安全上限爬取最多可访问的页面")
        self._report('start', {'url': self.base_url, 'max_pages': max_pages_int})
        
        # 获取网站标题
        try:
            content, soup = self.get_page_content(self.base_url)
            if soup:
                title = soup.find('title')
                if title:
                    self.site_title = self.clean_text(title.get_text().strip())
                    logger.info(f"网站标题: {self.site_title}")
                    self._report('site_title', {'title': self.site_title})
        except Exception as e:
            logger.warning(f"获取网站标题失败: {e}")
        
        # 安全机制：最大爬取10000页，避免无限爬取
        safety_limit = 10000
        
        while urls_to_visit:
            current_url = urls_to_visit.pop(0)
            
            if current_url in self.visited_urls:
                continue
            
            # 安全限制检查
            if page_count >= safety_limit:
                logger.warning(f"已达到安全限制 {safety_limit} 页，停止爬取")
                logger.warning("如需继续，请修改代码中的 safety_limit 值")
                break
            # 页数限制检查
            if max_pages_int is not None and page_count >= max_pages_int:
                logger.info(f"已达到设定的最大页数 {max_pages_int}，停止爬取")
                break
                
            logger.info(f"正在爬取第 {page_count + 1} 页: {current_url}")
            self._report('page_start', {'index': page_count + 1, 'url': current_url, 'queue': len(urls_to_visit)})
            
            # 获取页面内容
            content, soup = self.get_page_content(current_url)
            if not soup:
                continue
            
            # 标记为已访问
            self.visited_urls.add(current_url)
            page_count += 1
            
            # 提取页面信息
            before_phones = len(self.seen_phones)
            before_contacts = len(self.seen_contacts)
            self.extract_page_info(current_url, soup)
            delta_phones = len(self.seen_phones) - before_phones
            delta_contacts = len(self.seen_contacts) - before_contacts
            self._report('page_result', {
                'index': page_count,
                'url': current_url,
                'new_phones': max(0, delta_phones),
                'new_contacts': max(0, delta_contacts)
            })
            
            # 查找新链接
            new_links = self.find_all_links(soup, current_url)
            for link in new_links:
                if link not in self.visited_urls and link not in urls_to_visit:
                    urls_to_visit.append(link)
            
            # 显示进度信息
            if page_count % 10 == 0:
                logger.info(f"爬取进度: 已爬取 {page_count} 页，待爬取 {len(urls_to_visit)} 页")
                logger.info(f"已找到 {len(self.seen_phones)} 个手机号，{len(self.seen_contacts)} 个联系人")
                self._report('progress', {
                    'pages': page_count,
                    'queue': len(urls_to_visit),
                    'phones': len(self.seen_phones),
                    'contacts': len(self.seen_contacts)
                })
            
            # 避免请求过快
            time.sleep(1)
        
        logger.info(f"爬取完成，共爬取 {page_count} 页")
        logger.info(f"总计 {len(self.seen_phones)} 个手机号，{len(self.seen_contacts)} 个联系人")
        self._report('done', {
            'pages': page_count,
            'phones': len(self.seen_phones),
            'contacts': len(self.seen_contacts)
        })
    
    def export_to_csv(self, filename: str = 'phone_contacts.csv') -> None:
        """导出结果到CSV文件"""
        if not self.phone_contacts:
            logger.warning("没有找到任何手机号码或联系人信息")
            return
            
        try:
            with open(filename, 'w', newline='', encoding='utf-8-sig') as csvfile:
                fieldnames = ['url', 'title', 'phone_numbers', 'contacts', 'phone_count', 'contact_count']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                
                writer.writeheader()
                for row in self.phone_contacts:
                    safe_row = {key: row.get(key, '') for key in fieldnames}
                    writer.writerow(safe_row)
            
            logger.info(f"结果已导出到 {filename}")
            logger.info(f"共找到 {len(self.phone_contacts)} 页包含联系信息")
            
            # 统计总手机号和联系人数量
            total_phones = sum(row['phone_count'] for row in self.phone_contacts)
            total_contacts = sum(row['contact_count'] for row in self.phone_contacts)
            logger.info(f"总计: {total_phones} 个手机号, {total_contacts} 个联系人")
            
        except Exception as e:
            logger.error(f"导出CSV失败: {e}")
    
    def export_to_json(self, filename: str = 'phone_contacts.json') -> None:
        """导出结果到JSON文件"""
        if not self.phone_contacts:
            return
            
        try:
            with open(filename, 'w', encoding='utf-8') as jsonfile:
                json.dump(self.phone_contacts, jsonfile, ensure_ascii=False, indent=2)
            
            logger.info(f"结果已导出到 {filename}")
        except Exception as e:
            logger.error(f"导出JSON失败: {e}")
    
    def export_to_docx(self, filename: str = 'phone_contacts.docx') -> None:
        """导出结果到Word文档"""
        if not self.phone_contacts:
            logger.warning("没有找到任何手机号码或联系人信息")
            return
            
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置标题
            title = doc.add_heading(self.site_title, 0)
            title.alignment = 1  # 居中对齐
            
            # 添加说明
            doc.add_paragraph(f"爬取时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"爬取页面数: {len(self.phone_contacts)}")
            doc.add_paragraph(f"总计手机号: {len(self.seen_phones)} 个")
            doc.add_paragraph(f"总计联系人: {len(self.seen_contacts)} 个")
            
            # 添加分隔线
            doc.add_paragraph("=" * 50)
            
            # 添加内容
            for i, result in enumerate(self.phone_contacts, 1):
                # 页面标题
                page_heading = doc.add_heading(f"{i}. {result['title']}", level=1)
                
                # 手机号信息
                if result['phone_numbers']:
                    phone_para = doc.add_paragraph()
                    phone_para.add_run("手机号: ").bold = True
                    phone_para.add_run(result['phone_numbers'])
                
                # 联系人信息
                if result['contacts']:
                    contact_para = doc.add_paragraph()
                    contact_para.add_run("联系人: ").bold = True
                    contact_para.add_run(result['contacts'])
                
                # 页面URL（小字）
                url_para = doc.add_paragraph()
                url_run = url_para.add_run(f"来源: {result['url']}")
                url_run.font.size = doc.styles['Normal'].font.size
                url_run.font.color.rgb = None  # 默认颜色
                
                # 添加分隔线
                doc.add_paragraph("-" * 30)
            
            # 保存文档
            doc.save(filename)
            logger.info(f"结果已导出到Word文档: {filename}")
            
        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")

def main():
    """主函数"""
    # 目标网站
    target_url = "https://www.schdri.com/go.htm?k=zhong_dian_gong_cheng&url=cheng_guo_zhan_shi/zhong_dian_gong_cheng"
    
    # 创建爬虫实例
    scraper = PhoneScraper(target_url)
    
    try:
        # 开始爬取 - 无限制模式
        scraper.crawl_website()
        
        # 导出结果
        scraper.export_to_csv('phone_contacts.csv')
        scraper.export_to_json('phone_contacts.json')
        scraper.export_to_docx('phone_contacts.docx')  # 新增Word导出
        
        # 显示结果摘要
        print("\n" + "="*50)
        print("爬取结果摘要:")
        print("="*50)
        
        for i, info in enumerate(scraper.phone_contacts[:10], 1):  # 显示前10页
            print(f"\n{i}. {info['title']}")
            print(f"   URL: {info['url']}")
            if info['phone_numbers']:
                print(f"   手机号: {info['phone_numbers']}")
            if info['contacts']:
                print(f"   联系人: {info['contacts']}")
        
        if len(scraper.phone_contacts) > 10:
            print(f"\n... 还有 {len(scraper.phone_contacts) - 10} 页结果")
        
        print(f"\n完整结果已导出到 phone_contacts.csv 和 phone_contacts.json")
        
    except KeyboardInterrupt:
        logger.info("用户中断爬取")
        # 即使中断也导出已爬取的结果
        if scraper.phone_contacts:
            scraper.export_to_csv('phone_contacts_partial.csv')
            scraper.export_to_json('phone_contacts_partial.json')
            print("已导出部分结果")
    except Exception as e:
        logger.error(f"爬取过程中发生错误: {e}")
        # 导出已爬取的结果
        if scraper.phone_contacts:
            scraper.export_to_csv('phone_contacts_error.csv')
            scraper.export_to_json('phone_contacts_error.json')

if __name__ == "__main__":
    main() 