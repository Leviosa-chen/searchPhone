#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试Word导出功能
"""

from docx import Document
from docx.shared import Inches
import time

def test_word_export():
    """测试Word导出功能"""
    print("=" * 60)
    print("测试Word导出功能")
    print("=" * 60)
    
    # 模拟爬取结果
    test_results = [
        {
            'url': 'https://example.com/page1',
            'title': '张三的联系方式',
            'phones': ['13800138000', '13900139000'],
            'contacts': ['张三', '李四']
        },
        {
            'url': 'https://example.com/page2',
            'title': '王五的联系方式',
            'phones': ['13600136000'],
            'contacts': ['王五']
        },
        {
            'url': 'https://example.com/page3',
            'title': '赵六的联系方式',
            'phones': ['13500135000'],
            'contacts': ['赵六', '钱七']
        }
    ]
    
    print("模拟数据:")
    for i, result in enumerate(test_results, 1):
        print(f"{i}. {result['title']}")
        print(f"   手机号: {', '.join(result['phones'])}")
        print(f"   联系人: {', '.join(result['contacts'])}")
        print()
    
    # 创建Word文档
    print("正在创建Word文档...")
    doc = Document()
    
    # 设置标题
    site_title = "示例网站 - 联系方式汇总"
    title = doc.add_heading(site_title, 0)
    title.alignment = 1  # 居中对齐
    
    # 添加说明
    doc.add_paragraph(f"爬取时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"爬取页面数: {len(test_results)}")
    
    # 统计总数
    total_phones = sum(len(r['phones']) for r in test_results)
    total_contacts = sum(len(r['contacts']) for r in test_results)
    doc.add_paragraph(f"总计手机号: {total_phones} 个")
    doc.add_paragraph(f"总计联系人: {total_contacts} 个")
    
    # 添加分隔线
    doc.add_paragraph("=" * 50)
    
    # 添加内容
    for i, result in enumerate(test_results, 1):
        # 页面标题
        page_heading = doc.add_heading(f"{i}. {result['title']}", level=1)
        
        # 手机号信息
        if result['phones']:
            phone_para = doc.add_paragraph()
            phone_para.add_run("手机号: ").bold = True
            phone_para.add_run(", ".join(result['phones']))
        
        # 联系人信息
        if result['contacts']:
            contact_para = doc.add_paragraph()
            contact_para.add_run("联系人: ").bold = True
            contact_para.add_run(", ".join(result['contacts']))
        
        # 页面URL（小字）
        url_para = doc.add_paragraph()
        url_run = url_para.add_run(f"来源: {result['url']}")
        url_run.font.size = doc.styles['Normal'].font.size
        url_run.font.color.rgb = None  # 默认颜色
        
        # 添加分隔线
        doc.add_paragraph("-" * 30)
    
    # 保存文档
    filename = 'test_export.docx'
    doc.save(filename)
    
    print(f"✓ Word文档创建成功: {filename}")
    print(f"✓ 文档包含 {len(test_results)} 个页面")
    print(f"✓ 总计 {total_phones} 个手机号, {total_contacts} 个联系人")
    
    return filename

def test_text_cleaning():
    """测试文本清理功能"""
    print("\n" + "=" * 60)
    print("测试文本清理功能")
    print("=" * 60)
    
    import re
    
    def clean_text(text):
        """清理文本，去掉特殊字符"""
        if not text:
            return ""
        
        # 去掉特殊字符，保留中文、英文、数字、空格、冒号、括号
        cleaned = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s：:()（）]', '', text)
        # 去掉多余空格
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned
    
    # 测试用例
    test_cases = [
        "张三的联系方式 - 手机号：13800138000",
        "李四 (经理) 电话: 13900139000",
        "王五@example.com 手机：13600136000",
        "赵六 & 钱七 联系方式：13500135000",
        "孙八【技术部】电话：13400134000",
        "周九 <hr> 手机：13300133000"
    ]
    
    print("原始文本 -> 清理后:")
    for i, text in enumerate(test_cases, 1):
        cleaned = clean_text(text)
        print(f"{i}. {text}")
        print(f"   -> {cleaned}")
        print()

def main():
    """主函数"""
    print("开始测试Word导出功能...")
    
    # 测试文本清理
    test_text_cleaning()
    
    # 测试Word导出
    filename = test_word_export()
    
    print("\n" + "=" * 60)
    print("测试完成！")
    print("=" * 60)
    print(f"\n生成的Word文档: {filename}")
    print("请打开查看效果")
    print("\n实际使用中，爬虫会自动:")
    print("1. 清理页面标题中的特殊字符")
    print("2. 只导出包含联系信息的页面")
    print("3. 生成格式化的Word文档")

if __name__ == "__main__":
    main() 