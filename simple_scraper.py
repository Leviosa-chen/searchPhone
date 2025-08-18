#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简化版网站手机号码和联系人爬虫
专门用于爬取 https://www.schdri.com 网站
"""

import requests
from bs4 import BeautifulSoup
import re
import csv
import time
import json
from urllib.parse import urljoin, urlparse
from docx import Document
from docx.shared import Inches
import os

class SimplePhoneScraper:
    def __init__(self):
        self.base_url = "https://www.schdri.com"
        self.start_url = "https://www.schdri.com/go.htm?k=zhong_dian_gong_cheng&url=cheng_guo_zhan_shi/zhong_dian_gong_cheng"
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        })
        self.visited = set()
        self.results = []
        # 用于跟踪已出现的手机号，确保不重复
        self.seen_phones = set()
        # 页面层级跟踪
        self.page_levels = {}  # 记录每个页面的层级
        
    def clean_text(self, text):
        """清理文本，去掉特殊字符"""
        if not text:
            return ""
        
        # 去掉特殊字符，保留中文、英文、数字、空格、冒号、括号
        cleaned = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s：:()（）]', '', text)
        # 去掉多余空格
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        return cleaned
    
    def extract_phones(self, text):
        """提取手机号码 - 优化版，确保前后不是数字"""
        # 中国手机号码模式 - 使用负向前瞻和负向后瞻确保前后不是数字
        patterns = [
            r'(?<!\d)1[3-9]\d{9}(?!\d)',  # 标准11位，前后不能是数字
            r'(?<!\d)\+86\s*1[3-9]\d{9}(?!\d)',  # 带+86前缀
            r'(?<!\d)86\s*1[3-9]\d{9}(?!\d)',  # 带86前缀
            r'(?<!\d)1[3-9]\d{2}\s*\d{4}\s*\d{4}(?!\d)',  # 带空格的手机号
            r'(?<!\d)1[3-9]\d{2}-\d{4}-\d{4}(?!\d)',  # 带连字符的手机号
        ]
        
        phones = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                # 清理号码格式
                clean_phone = re.sub(r'[\s\+\-]', '', match)
                if len(clean_phone) == 11 and clean_phone.startswith('1'):
                    # 额外验证：确保不是文件名的一部分
                    if not self._is_filename_part(clean_phone, text):
                        phones.append(clean_phone)
        
        return list(set(phones))  # 去重
    
    def _is_filename_part(self, phone, text):
        """检查手机号是否是文件名的一部分"""
        # 查找手机号在文本中的位置
        phone_pos = text.find(phone)
        if phone_pos == -1:
            return False
        
        # 检查前后字符，判断是否是文件名
        start_pos = max(0, phone_pos - 10)
        end_pos = min(len(text), phone_pos + 21)
        context = text[start_pos:end_pos]
        
        # 如果包含常见文件扩展名，可能是文件名
        file_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.pdf', '.doc', '.docx', '.xls', '.xlsx']
        for ext in file_extensions:
            if ext in context.lower():
                return True
        
        # 如果前后都是数字，可能是长数字的一部分
        if phone_pos > 0 and phone_pos + 11 < len(text):
            prev_char = text[phone_pos - 1]
            next_char = text[phone_pos + 11]
            if prev_char.isdigit() and next_char.isdigit():
                return True
        
        return False
    
    def extract_contacts(self, text):
        """提取联系人 - 优化版，注意前后语义"""
        # 改进的联系人正则表达式，更注重语义
        contact_patterns = [
            # 标准格式：关键词: 联系人
            r'(?:联系人|负责人|经理|主管|主任|总监|姓名|名字)[：:]\s*([^\n\r]{1,15})',
            
            # 带职务的格式：联系人：张三 经理
            r'(?:联系人|负责人|经理|主管|主任|总监)[：:]\s*([^\n\r]{1,20})',
            
            # 表格格式：联系人 张三
            r'(?:联系人|负责人|经理|主管|主任|总监)\s+([^\n\r]{1,15})',
            
            # 姓名格式：姓名：张三
            r'姓名[：:]\s*([^\n\r]{1,15})',
        ]
        
        contacts = []
        for pattern in contact_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                contact = match.strip()
                if self._is_valid_contact(contact):
                    contacts.append(contact)
        
        return list(set(contacts))  # 去重
    
    def _is_valid_contact(self, contact):
        """验证联系人信息是否有效"""
        if not contact:
            return False
        
        # 长度检查
        if len(contact) < 2 or len(contact) > 20:
            return False
        
        # 排除明显无效的内容
        invalid_patterns = [
            r'^\d+$',  # 纯数字
            r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$',  # 邮箱
            r'^1[3-9]\d{9}$',  # 手机号
            r'^\d{3,4}-\d{7,8}$',  # 座机号
            r'^[^\u4e00-\u9fa5a-zA-Z]*$',  # 不包含中文或英文
        ]
        
        for pattern in invalid_patterns:
            if re.match(pattern, contact):
                return False
        
        # 必须包含中文或英文
        if not re.search(r'[\u4e00-\u9fa5a-zA-Z]', contact):
            return False
        
        return True
    
    def get_page(self, url):
        """获取页面内容"""
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            response.encoding = 'utf-8'
            return response.text
        except Exception as e:
            print(f"获取页面失败 {url}: {e}")
            return None
    
    def find_links(self, html, current_url):
        """查找页面中的链接"""
        if not html:
            return []
        
        soup = BeautifulSoup(html, 'html.parser')
        links = []
        
        for a in soup.find_all('a', href=True):
            href = a['href']
            full_url = urljoin(current_url, href)
            
            # 只保留同域名的链接
            if urlparse(full_url).netloc == 'www.schdri.com':
                links.append(full_url)
        
        return links
    
    def crawl(self, max_pages=1000, max_level=6):
        """开始爬取 - 限制1000页，最多6级页面"""
        urls_to_visit = [(self.start_url, 0)]  # (url, level)
        page_count = 0
        self.site_title = "未知网站"  # 默认网站标题
        
        print(f"开始爬取网站: {self.start_url}")
        print("=" * 60)
        print(f"爬取限制: 最多{max_pages}页，最多{max_level}级页面")
        print("爬取内容: 仅手机号（不爬取联系人）")
        print("如需停止，请按 Ctrl+C")
        print("=" * 60)
        
        # 获取网站标题
        first_page = self.get_page(self.start_url)
        if first_page:
            soup = BeautifulSoup(first_page, 'html.parser')
            title = soup.find('title')
            if title:
                self.site_title = self.clean_text(title.get_text().strip())
                print(f"网站标题: {self.site_title}")
        
        while urls_to_visit and page_count < max_pages:
            current_url, current_level = urls_to_visit.pop(0)
            
            if current_url in self.visited:
                continue
            
            # 检查页面层级
            if current_level > max_level:
                print(f"⚠️  跳过页面 {current_url} - 层级 {current_level} 超过限制 {max_level}")
                continue
            
            print(f"正在爬取第 {page_count + 1} 页 (层级 {current_level}): {current_url}")
            
            # 获取页面内容
            html = self.get_page(current_url)
            if not html:
                continue
            
            self.visited.add(current_url)
            self.page_levels[current_url] = current_level
            page_count += 1
            
            # 只提取手机号，不提取联系人
            phones = self.extract_phones(html)
            
            # 去重处理：只保留首次出现的手机号
            unique_phones = []
            for phone in phones:
                if phone not in self.seen_phones:
                    unique_phones.append(phone)
                    self.seen_phones.add(phone)
            
            # 获取页面标题
            soup = BeautifulSoup(html, 'html.parser')
            title = soup.find('title')
            page_title = self.clean_text(title.get_text().strip()) if title else "无标题"
            
            # 记录结果 - 只记录有手机号的页面
            if unique_phones:
                result = {
                    'url': current_url,
                    'title': page_title,
                    'phones': unique_phones,
                    'phone_count': len(unique_phones),
                    'level': current_level,
                    'original_phones': phones  # 保留原始数据用于对比
                }
                self.results.append(result)
                
                print(f"  ✓ 找到 {len(unique_phones)} 个新手机号")
                if unique_phones:
                    print(f"     新手机号: {', '.join(unique_phones)}")
                
                # 如果有重复数据，显示去重信息
                duplicate_phones = len(phones) - len(unique_phones)
                if duplicate_phones > 0:
                    print(f"     (去重: {duplicate_phones} 个重复手机号)")
            else:
                print(f"  - 未找到新的手机号")
            
            # 查找新链接，但限制层级
            if current_level < max_level:
                new_links = self.find_links(html, current_url)
                for link in new_links:
                    if link not in self.visited and not any(link == url for url, _ in urls_to_visit):
                        urls_to_visit.append((link, current_level + 1))
            
            # 显示进度信息
            if page_count % 10 == 0:
                print(f"\n📊 爬取进度: 已爬取 {page_count}/{max_pages} 页，待爬取 {len(urls_to_visit)} 页")
                print(f"📱 已找到 {len(self.seen_phones)} 个手机号")
                print(f"🌐 当前最高层级: {max(self.page_levels.values()) if self.page_levels else 0}")
            
            # 控制爬取速度
            time.sleep(0.5)
        
        print(f"\n爬取完成！共爬取 {page_count} 页")
        print(f"找到 {len(self.results)} 页包含手机号")
        print(f"总计 {len(self.seen_phones)} 个手机号")
        print(f"最高页面层级: {max(self.page_levels.values()) if self.page_levels else 0}")
        
        # 显示层级统计
        level_stats = {}
        for level in self.page_levels.values():
            level_stats[level] = level_stats.get(level, 0) + 1
        
        print(f"\n页面层级分布:")
        for level in sorted(level_stats.keys()):
            print(f"  层级 {level}: {level_stats[level]} 页")
    
    def export_docx(self, filename='phone_contacts.docx'):
        """导出到Word文档 - 仅手机号"""
        if not self.results:
            print("没有找到任何手机号信息")
            return
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置标题
            title = doc.add_heading(self.site_title, 0)
            title.alignment = 1  # 居中对齐
            
            # 添加说明
            doc.add_paragraph(f"爬取时间: {time.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"爬取页面数: {len(self.results)}")
            doc.add_paragraph(f"总计手机号: {len(self.seen_phones)} 个")
            doc.add_paragraph(f"最高页面层级: {max(self.page_levels.values()) if self.page_levels else 0}")
            
            # 添加分隔线
            doc.add_paragraph("=" * 50)
            
            # 添加内容
            for i, result in enumerate(self.results, 1):
                # 页面标题
                page_heading = doc.add_heading(f"{i}. {result['title']}", level=1)
                
                # 手机号信息
                if result['phones']:
                    phone_para = doc.add_paragraph()
                    phone_para.add_run("手机号: ").bold = True
                    phone_para.add_run(", ".join(result['phones']))
                
                # 页面层级信息
                level_para = doc.add_paragraph()
                level_para.add_run("页面层级: ").bold = True
                level_para.add_run(f"第 {result['level']} 级")
                
                # 页面URL（小字）
                url_para = doc.add_paragraph()
                url_run = url_para.add_run(f"来源: {result['url']}")
                url_run.font.size = doc.styles['Normal'].font.size
                url_run.font.color.rgb = None  # 默认颜色
                
                # 添加分隔线
                doc.add_paragraph("-" * 30)
            
            # 保存文档
            doc.save(filename)
            print(f"\n结果已导出到Word文档: {filename}")
            
        except Exception as e:
            print(f"导出Word文档失败: {e}")
    
    def export_csv(self, filename='phone_contacts.csv'):
        """导出到CSV - 仅手机号"""
        if not self.results:
            print("没有找到任何手机号信息")
            return
        
        try:
            with open(filename, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                writer.writerow(['页面标题', 'URL', '手机号码', '手机号数量', '页面层级'])
                
                for result in self.results:
                    writer.writerow([
                        result['title'],
                        result['url'],
                        '; '.join(result['phones']),
                        result['phone_count'],
                        result['level']
                    ])
            
            print(f"\n结果已导出到: {filename}")
            
            # 统计总数
            total_phones = sum(r['phone_count'] for r in self.results)
            print(f"总计: {total_phones} 个手机号")
            
        except Exception as e:
            print(f"导出CSV失败: {e}")
    
    def export_json(self, filename='phone_contacts.json'):
        """导出到JSON"""
        if not self.results:
            return
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(self.results, f, ensure_ascii=False, indent=2)
            print(f"结果已导出到: {filename}")
        except Exception as e:
            print(f"导出JSON失败: {e}")

def main():
    """主函数"""
    scraper = SimplePhoneScraper()
    
    try:
        # 开始爬取 - 限制1000页，最多6级页面
        scraper.crawl(max_pages=1000, max_level=6)
        
        # 导出结果
        scraper.export_csv()
        scraper.export_json()
        scraper.export_docx() # 新增的导出方法
        
        # 显示结果摘要
        if scraper.results:
            print("\n" + "=" * 60)
            print("爬取结果摘要:")
            print("=" * 60)
            
            for i, result in enumerate(scraper.results[:5], 1):
                print(f"\n{i}. {result['title']}")
                print(f"   URL: {result['url']}")
                if result['phones']:
                    print(f"   手机号: {', '.join(result['phones'])}")
                if result['contacts']:
                    print(f"   联系人: {', '.join(result['contacts'])}")
            
            if len(scraper.results) > 5:
                print(f"\n... 还有 {len(scraper.results) - 5} 页结果")
        
    except KeyboardInterrupt:
        print("\n\n用户中断爬取")
        if scraper.results:
            scraper.export_csv('phone_contacts_partial.csv')
            scraper.export_json('phone_contacts_partial.json')
            print("已导出部分结果")
    except Exception as e:
        print(f"\n爬取过程中发生错误: {e}")
        if scraper.results:
            scraper.export_csv('phone_contacts_error.csv')
            scraper.export_json('phone_contacts_error.json')

if __name__ == "__main__":
    main() 